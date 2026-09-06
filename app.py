import os
import sys
import time
import json
import re
import random
import base64
import datetime
import tempfile
import subprocess
from pathlib import Path

import streamlit as st
import pypandoc
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from google import genai
from google.genai import types
from google.genai.errors import APIError

# ---------------------------------------------------------
# STREAMLIT PAGE CONFIGURATION
# ---------------------------------------------------------
st.set_page_config(
    page_title="iWish Exam Paper Generator",
    page_icon="🪄",
    layout="wide"
)

# ---------------------------------------------------------
# SECURE API KEY RESOLUTION
# ---------------------------------------------------------
def resolve_api_key() -> str:
    if "GEMINI_API_KEY" in st.secrets:
        return st.secrets["GEMINI_API_KEY"]
    if os.environ.get("GEMINI_API_KEY"):
        return os.environ.get("GEMINI_API_KEY")
    return ""

# ---------------------------------------------------------
# TOP 10 PRODUCTION MODEL FALLBACK HIERARCHY
# ---------------------------------------------------------
STATIC_MODELS_TO_TRY = [
    "gemini-3.8-flash",          # Rank 1: Primary (Flagship agentic reasoning & multimodal speed)
    "gemini-3.7-flash",          # Rank 2: Fallback 1 (High intelligence-to-speed ratio)
    "gemini-3.6-flash",          # Rank 3: Fallback 2 (Google recommended stable production tier)
    "gemini-3.1-pro-preview",    # Rank 4: Fallback 3 (Deep reasoning for complex calculations)
    "gemini-3.5-flash",          # Rank 5: Fallback 4 (Proven multimodal workhorse)
    "gemini-2.5-pro",            # Rank 6: Fallback 5 (Advanced structural reasoning)
    "gemini-flash-latest",       # Rank 7: Fallback 6 (Dynamic active Flash alias)
    "gemini-pro-latest",         # Rank 8: Fallback 7 (Dynamic active Pro alias)
    "gemini-3.5-flash-lite",     # Rank 9: Fallback 8 (Low-latency high-throughput tier)
    "gemini-3.1-flash-lite",     # Rank 10: Fallback 9 (Emergency speed fallback)
]

# ---------------------------------------------------------
# 300 HINGLISH "MAGIC IN PROGRESS" MESSAGES (50 PER TIER)
# ---------------------------------------------------------
TIER_1_MESSAGES = [
    "Mendeleev aur Newton ki aatma ko summon kiya jaa raha hai... 🧙‍♂️✨",
    "Chai ka cup utha lijiye, homework ka asli jaadu ab shuru ho raha hai... ☕🪄",
    "PDF ko AI ke hawaale kar diya hai, ab bas aage ka tamasha dekhiye... 🍿🚀",
    "Teacher mode: ON! Bachhon ki azaadi bas kuch der ki mehmaan hai... 😈⏰",
    "Shanti se baithiye, homework ka brahmastra tayyar ho raha hai... 🏹⚡",
    "System garam ho raha hai aur students ka future decide ho raha hai... 🔥💻",
    "Almirah se purane board papers aur secret notes nikaale jaa rahe hain... 📚🕵️‍♂️",
    "Backbenchers ke chain aur sukoon par strike karne ka time aa gaya hai... 🎯💣",
    "Homework engine start ho chuka hai, agle 2 minute mein paper ready... 🏎️💨",
    "Einstein ka dimaag aur Newton ka gussa compile kiya jaa raha hai... 🧠⚡",
    "Aaj ka homework aisa banega jo seedha dil pe nahi, marks pe lagega... 💔📝",
    "Toppers ke chehre ki smile gayab karne ka blueprint tayyar ho raha hai... 😈📄",
    "Coffee pee lijiye sir/ma'am, paper banana ab humari zimmedari hai... ☕💻",
    "Class ke sabse shant bacche ko bhi sochne pe majboor karne ki taiyari... 🤫🤯",
    "Ek aur weekend barbaad karne ka poora intezam shuru ho chuka hai... 🏖️❌",
    "Magic wand ghuma di hai, PDF ke andar se questions kheenche jaa rahe hain... 🪄📜",
    "Exam hall wala sannata homework sheet mein bharne ka plan chalu hai... 🥶⚡",
    "Bachhe sochenge teacher itne creative kab se ho gaye... 🧐🎨",
    "Question paper generation engine 1000 RPM pe daud raha hai... ⚙️🚀",
    "Sunday ko fun day nahi, study day banane ka shubh aarambh... 🗓️✍️",
    "Aaj homework dekh kar WhatsApp class group mein dange hone wale hain... 📱🔥",
    "Science ke saare musibat bhare formulas ek jagah jama ho rahe hain... 🧬🧪",
    "Syllabus ke sabse uninteresting topics ko spicy banane ki taiyari... 🌶️📖",
    "Google Lens ko bhi chakkar aa jaye, aisa base ready ho raha hai... 🌀📱",
    "Shuru ho chuka hai mission: 'Homework dekh ke paseene chhootenge'... 😅💦",
    "Pen aur notebook uthane ka waqt aa gaya hai, dosto... 📝🎒",
    "Bachhon ke phone usage time ko zero karne ka secret weapon loading... 📵🛡️",
    "Laboratory ke sabse explosive concepts ka digital blast hone wala hai... 💥🔬",
    "Aaj ki sham bachhon ki homework copy ke naam... 🕯️📖",
    "Sawaal itne shaandar honge ki external examiner bhi salaam thoke... 🫡📄",
    "Aisa paper ban raha hai jise dekh ke viva ka darr feeka pad jaye... 🥶🎤",
    "Background mein magic run ho raha hai, front pe sab calm hai... 🌊🧙‍♂️",
    "AI examiner ne chashma laga liya hai aur laal pen utha liya hai... 👓🔴",
    "Students ke bahane khatam karne ka fool-proof system load ho raha hai... 🛑📑",
    "Aaj syllabus ke chhupe hue raaz bahar niklenge... 🗝️📂",
    "Homework ka chakravyuh ban raha hai, Abhimanyu kaun banega dekhte hain... 🔄🏹",
    "Bachhon ki neend churane ka legal software active ho chuka hai... 😴🚫",
    "Thoda sabr rakhiye, ek masterclass assessment shape le raha hai... ⏳💎",
    "Aaj ratta maarne walo ki calculation hilne wali hai... 📉🤯",
    "Class monitor bhi do baar sochega ki yeh kya de diya... 🧑‍💼❓",
    "Padhai ka maha-yuddh shuru karne ka conch blow ho chuka hai... 🐚⚔️",
    "AI server pe chemistry ke formulas garma-garam pak rahe hain... 🍲⚗️",
    "Bachhon ki reels dekhne ki aadat par full stop lagane ka process chalu... 🛑📱",
    "Masterstroke homework plan loading, seatbelt baandh lijiye... 🛫💺",
    "Hard work aur homework ka deadly combination taiyar ho raha hai... 💀📚",
    "Har sawaal mein ek naya jhatka dene ki ninja technique load ho rahi hai... 🥷⚡",
    "Har page par dimaag kholne wali exercise taiyar ki jaa rahi hai... 🏋️‍♂️🧠",
    "Question bank ki tijori khol di gayi hai... 🏦🔑",
    "Homework checking ke time aane wali muskaan pehle se book kar lijiye... 😏🖊️",
    "Loading completed partially, abhi to party shuru hui hai... 🎉🚀"
]

TIER_2_MESSAGES = [
    "Aapki PDF aise padhi jaa rahi hai jaise exam ki raat 3 baje padhai hoti hai... 📖⚡",
    "AI poori PDF ko ghol ke pee raha hai, ek-ek concept dhoondh raha hai... 🧪🔍",
    "Topper ki tarah line-by-line scanning chalu hai taaki kuch na chhoote... 🧐📄",
    "PDF se sabse khatarnak aur tedious topics dhoondh ke nikaale jaa rahe hain... ⚠️🔬",
    "Scanner mode ON: Har diagram, chart aur footnote ko scan kiya jaa raha hai... 📊👀",
    "PDF ke font size 8 wale chhupe hue notes bhi pakad liye gaye hain... 🔍🧐",
    "Textbook ke 'Do You Know?' wale corner se sawaal dhoondhe jaa rahe hain... 💡📚",
    "Ek-ek sentence ka post-mortem chal raha hai... 🩺📄",
    "Itni tez reading ho rahi hai ki PDF ke pixels bhi garam ho gaye hain... ⚡🔥",
    "Chapter ka summary nahi, andar ka core masala extract kiya jaa raha hai... 🌶️📖",
    "Har page ki aisi scanning ho rahi hai jaise board ki checking team karti hai... 🕵️‍♂️📑",
    "Exception wale cases ko khas tor par spotlight mein laya jaa raha hai... 🔦🪤",
    "Book ke back exercises se bhi zyada twist wale lines pakad rahe hain... 🥨📚",
    "PDF padhte padhte AI khud keh raha hai: 'Waah, kya theory hai!'... 🤖👌",
    "Speed aisi hai ki 50 page ka chapter 5 second mein digest ho gaya... 🌪️📄",
    "Chhote se chhote definition aur derivation ka x-ray scan ho raha hai... 🩻🔬",
    "Jo lines bachhe skip kar dete hain, wahi se questions extract ho rahe hain... 🎯🤫",
    "PDF ka ek-ek page bol raha hai: 'Mujhe mat chhedo!'... 📜💥",
    "Concept ke deep ocean mein gota lagaya jaa raha hai... 🌊🤿",
    "High-yield conceptual gold mines dhoondhi jaa rahi hain PDF ke andar... ⛏️🪙",
    "Speed reading level: Infinite! Human teacher bhi sharma jaye... 🚀📖",
    "Syllabus ke andhere kono mein torch maar ke sawaal dhoondh rahe hain... 🔦🕳️",
    "'Is paragraph se kya sawaal ban sakta hai?' investigation chal rahi hai... 🔍🤔",
    "PDF ke har formula ko test-tube mein daal ke shake kiya jaa raha hai... 🧪🌪️",
    "Aisa analysis chal raha hai jaise forensic lab mein investigation ho... 🔬🥼",
    "Bachhon ke 'Sir yeh book mein nahi tha' wale excuse ki pehle hi checking... 🚫🤷",
    "Har ek sub-topic ki kundli nikali jaa rahi hai... 📜🔮",
    "Key terms aur scientific jargons ko filter paper se chhana jaa raha hai... ☕🧫",
    "Theory itni gehri scan ho rahi hai ki concepts surface pe tairne lage hain... 🏊‍♂️📚",
    "Important theorems ko highlight karke alag bucket mein daal rahe hain... 🪣💡",
    "Aisa lag raha hai jaise Sherlock Holmes PDF ki enquiry kar raha ho... 🕵️‍♂️🔎",
    "Saare tricky numericals ki list taiyar ho rahi hai background mein... 🔢📝",
    "Bachhe jo line pencil se underline nahi karte, wahi target pe hai... 🎯✏️",
    "Chapter ke end points aur starting notes ka connection joda jaa raha hai... 🔗📄",
    "Text parsing superfast chal rahi hai, CPU fan tez ghoom raha hai... 💻💨",
    "Theory ke beech mein chhupe hue numerical clues dhoondhe jaa rahe hain... 🧩💡",
    "Har chemical bond aur reaction equation ka inspection report taiyar... 📑🔬",
    "'Out of syllabus' ka koi chance nahi, sab book ke andar se hi hai... 📚✅",
    "Reading speed dekh ke book publisher bhi salute karega... 🫡📖",
    "Har definition ke peeche ka hidden logic extract ho raha hai... 🧠⚙️",
    "Graph, slope aur intercept ko millimeter accuracy ke sath padha jaa raha hai... 📈📐",
    "Bachhon ke guess work ki range se bahar ke points select ho rahe hain... 🏹🚫",
    "Concept filter 99.9% accuracy pe operate kar raha hai... 🛡️📊",
    "Aise sawaal nikal rahe hain jo sidha concepts ko hit karein... 🎯💥",
    "Har derivation ka first aur last step verify ho raha hai... 🪜🔢",
    "PDF parsing almost complete, saare gems bahar aa chuke hain... 💎📂",
    "Content extraction ka kaam topper ke photocopy notes se bhi tez hai... ⚡📋",
    "Har ek reaction arrow ka direction cross-check ho raha hai... ➡️🔄",
    "Raw data process ho chuka hai, ab manufacturing unit shuru hone wali hai... 🏭⚙️",
    "PDF scanning done! Ab shuru hoga asli sawaalon ka dangal... 🤼‍♂️📄"
]

TIER_3_MESSAGES = [
    "Aise homework questions dhoondh rahe hain jise dekh ke bachhe Google karna bhool jayein... 🤯📚",
    "ChatGPT bhi do baar soche aur 'I cannot answer' keh de, aise sawaal ban rahe hain... 🤖💥",
    "Backbenchers aur toppers dono ke dimaag ki batti gul karne ki taiyaari... 💡😵",
    "Numerical aise set ho rahe hain jisme calculator bhi haath khade kar de... 🧮🔥",
    "Copy-paste walo ke liye special surprise traps design ho rahe hain... 🪤❌",
    "'Bhai tu pehle solve kar, main copy karta hoon' wali dosti tootne ka time... 💔🤝",
    "Sawaal itne conceptual hain ki ratta maarne walo ka dimaag 404 Error dega... 🚫🧠",
    "Reaction mechanisms mein aise twist dale hain jaise suspense movie ka climax... 🎬🧪",
    "Sawaal chhota dikhega lekin solve karne mein 3 rough page bhar jayenge... 📄✍️",
    "'Yeh toh sir ne kabhi padhaya hi nahi' bolne ka mauka zero rakha jaa raha hai... 😇📚",
    "Pure logic-based questions ban rahe hain, tukka lagane ka koi scope nahi... 🎯🙅‍♂️",
    "Har sawaal mein ek silent speed-breaker lagaya jaa raha hai... 🚧🛑",
    "Calculation aisi rakhi hai ki decimal ke baad 3 digit tak paseena niklega... 🔢💦",
    "Bachhon ke study group mein emergency meeting bulane wale sawaal... 👥🚨",
    "Direct formula apply karne walo ke liye khas theoretical jhatka... ⚡📉",
    "Assertion aur Reason ke beech aisi jung chhidne wali hai ki maza aa jayega... ⚔️🤔",
    "Ek sawaal, teen concepts, aur dimaag ka complete breakdown... 💥🧠",
    "Short answer itne tight hain ki 2 line mein poori science mangenge... 📏✍️",
    "Long answer questions dekh ke bachhe sochenge 'Isse accha chhutti le lete'... 🏖️🤦‍♂️",
    "Questions mein aisi language use ho rahi hai ki samajhne mein 5 minute lagein... 🗣️⏳",
    "Brain workout questions compile ho rahe hain, gym jaane ki zaroorat nahi... 🏋️‍♂️💡",
    "Units ko itna smartly badla hai ki Joule aur Calorie mein danga ho jaye... 🔄⚖️",
    "Thermochemistry aur equilibrium ka aisa sangam jo pehle kabhi nahi dekha gaya... 🌪️🧪",
    "Direct seedhe sawaal humare dictionary mein nahi hain, twist zaroori hai... 🌀📖",
    "Sawaal simple dikhega, par ghuma ke poora syllabus touch karega... 🔄🧭",
    "Yeh homework dekh kar tuition teacher bhi solution dhoondhne lagega... 👨‍🏫📞",
    "'Dimaag ki dahi' karne wale elements ka perfect chemical balance ho gaya... 🥛🍋",
    "Thermodynamics ke entropy wale sawaal se bachhon ki entropy badhne wali hai... 📈🔥",
    "Concept clear hoga tabhi pen chalega, warna page kora hi rahega... ⚪🖊️",
    "Har question ek riddle ki tarah prepare kiya jaa raha hai... 🧩🧐",
    "Question papers ke history ka sabse interesting set ban raha hai... 📜✨",
    "Bachhe sochenge sir ne paper banaya hai ya CID ne case file... 🕵️‍♂️📁",
    "Question number 3 dekh kar room temperature pe paseena aana guaranteed hai... 🌡️😰",
    "Standard state aur standard conditions ke traps actively set kiye jaa rahe hain... 🏷️🪤",
    "Formula yaad hoga par apply kaise karein, yeh sochne mein time nikal jayega... ⏰🤯",
    "Creative problem solving skills ko test karne ka ultimate blueprint ban raha hai... 🎨🧪",
    "Har sawaal mein ek subtle twist jo sirf attentive bacha hi pakad payega... 🧠✨",
    "'Iska answer 0 ya 1 hoga' sochne walo ke liye khas boundary values... 0️⃣1️⃣",
    "Step-by-step marks scoring rubric ke hisab se structure tayyar ho raha hai... 🪜💯",
    "Theory ko real-life examples mein wrap karke present kar rahe hain... 🎁🌍",
    "Sawaal padhte hi bachhon ke dimaag ke saare neuron fire karne lagenge... ⚡🧠",
    "Sawaal aise hain ki toppers bhi do baar calculation recheck karenge... 🔁🧐",
    "No boring repetitive questions, har question mein ek naya challenge... 🆕🔥",
    "Question design complete hone ko hai, level ekdum standard... 🌍🏆",
    "Reaction pathways ko maze runner jaisa bana diya gaya hai... 🏃‍♂️🌀",
    "Har line mein ek academic punch chupaya gaya hai... 🥊📝",
    "Bachhe sochenge paper kisi professor ne banaya hai ya school teacher ne... 🏛️👨‍🎓",
    "High-order thinking skills (HOTS) ka maximum dosage load ho raha hai... 💉🧠",
    "Questions ki quality 24 carat gold se bhi zyada pure nikli hai... 🏅📄",
    "Drafting complete! Sawaal ekdum solid hain, ab options ki baari... 🔥🎯"
]

TIER_4_MESSAGES = [
    "Option B aur C mein thoda sa dimaag ka dahi karne wala twist daal rahe hain... 😈🎯",
    "Saare options ek jaise dikhein, aisi subtle ninja technique lagayi jaa rahi hai... 🥷✨",
    "'All of the above' daalein ya 'None of these'? Confuse karne ka formula set ho raha hai... 🎭🎲",
    "Option elimination method ki dhajjiyan udane wala kaam chal raha hai... 🧠🌪️",
    "Aisa option banaya hai jo pehli nazar mein 100% sahi lagega par hoga wrong... 🪤💔",
    "Plus (+) aur Minus (-) ka subtle khel jo har saal toppers ke marks khata hai... ➕➖",
    "Option A aisa rakha hai jo common calculation mistake ka exact result ho... 🎣🎣",
    "Charo options itne close hain ki lens laga kar dekhna padega... 🔬🔎",
    "'Jai Mata Di' bolke C mark karne walo ke plan pe paani phera jaa raha hai... 🎲🌊",
    "Units ki heera-pheri chalu hai: kJ mol⁻¹ ko J mol⁻¹ likh ke trap banaya... 🔄⚖️",
    "Option B bol raha hai 'Main sahi hoon', Option C bol raha hai 'Nahi, main hoon'... 🗣️😵",
    "Confident students ke overconfidence ko check karne wale options... 🧗‍♂️🪂",
    "Har MCQ option ek psychological game khelne ke liye design ho raha hai... ♟️🧠",
    "Negative marking ka ashirwad har galat option ke saath attach ho raha hai... 📉❌",
    "Ek option bilkul innocent dikhega, wahi sabse bada trap niklega... 😇🪤",
    "Four options, infinite confusion! Aisa algorithm active ho chuka hai... ♾️🌀",
    "Calculation karne ke baad jo sabse pehle galat number aayega, wahi Option A hai... 🎯😈",
    "Option D ko 'Both A and B' bana ke dimaag mein duvidha peda ki jaa rahi hai... ⚖️🤯",
    "Guess work karne walo ke luck ki battery low karne ka process chal raha hai... 🪫❌",
    "Har option ke font aur state symbols ko crystal clear rakha jaa raha hai... 👁️🔤",
    "Options dekh ke baccha bolega: 'Sir, options mein printing mistake hai kya?'... 🙋‍♂️🤦",
    "Distractor choices itne realistic hain ki book ke author bhi confuse ho jayein... 📚🤯",
    "Tukka lagane ki probability 25% se gira kar 0% ki jaa rahi hai... 📉🎲",
    "Scientific notation mein powers ke sath khelna chalu hai: 10⁶ vs 10⁻⁶... 🔢⚡",
    "Har option ka wording itna refined hai ki grammar se bhi clue na mile... 🔠🛡️",
    "Option C ko trap banaya hai kyunki 70% log C select karte hain... 📊🎯",
    "Baccha har option padhke bolega: 'Yeh to sab padha padha sa lag raha hai'... 😅📖",
    "Options ki ordering aisi hai ki pattern dhundhne walo ke pattern toot jayein... 📉🔨",
    "ABCD choices ab simple choices nahi, mind mazes ban chuki hain... 🌀🚪",
    "Ek option mein concept sahi hai par unit galat, dusre mein unit sahi concept galat... 🎭⚖️",
    "Topper bacha teen baar pencil se rough work mitayega... ✏️🧼",
    "Subtle changes in chemical formula: Fe²⁺ vs Fe³⁺ ka game set ho gaya... 🧪⚡",
    "Options ka balance aisa banaya hai ki answer key ekdum uniform lage... ⚖️📋",
    "Class test toppers ke 100/100 rokne wala ultimate barrier taiyar... 🚧💯",
    "Options ko decode karne ke liye pure dimaag ki 100% capacity lagegi... 🧠🔥",
    "Multiple choice questions ab multiple headache questions ban chuke hain... 🤕🎯",
    "'Almost correct' aur 'Perfectly correct' ke beech ka farq create kiya jaa raha hai... 🤏📏",
    "Options itne pyare hain ki tick karte waqt haath kaanpega... ✍️🥶",
    "Har wrong option ke peeche ek typical silly mistake chhuphi hui hai... 🙈🪤",
    "Baccha exam hall se nikal ke bolega: 'Paper easy tha par options khatarnak the'... 🗣️😩",
    "Option A, B, C, D mein confusion ka 4G network install kar diya gaya hai... 📶🌀",
    "Negative marking ka dar live stream hone wala hai... 🔴📉",
    "Options checking complete, distraction level maximum mark pe pahunch gaya... 🚀🎛️",
    "Intuition vs Logic ki ladai mein sirf Logic jeetega... 🥊💡",
    "Har option ko academic lab mein test karke deploy kiya jaa raha hai... 🔬🧪",
    "Elimination round chal raha hai, student ke dimaag ki kasrat tay hai... 🏋️‍♂️🧠",
    "Har option ek magnet ki tarah galat answer ki taraf kheenchega... 🧲🪤",
    "Options finalize ho chuke hain, ab paper ko formal dress pehnani hai... 👔📜",
    "MCQ block ekdum tight! Sawaal padh ke maza hi aa jayega... 🎯✨",
    "Options locked! Ab document ko stylish aur clean banate hain... 🔒🎨"
]

TIER_5_MESSAGES = [
    "Homework ko ekdum masoom look de rahe hain taaki pehle lage ki kitna aasan hai... 😇📝",
    "Page layout, clean fonts aur iWish logo ekdum aesthetic set ho rahe hain... 🎨📄",
    "Document formatting chal rahi hai, lagna chahiye board exam ka secret paper hai... 🤫📜",
    "Equations aur chemical formulas ko sundar sa LaTeX makeup lagaya jaa raha hai... 💄⚗️",
    "Arial font 11 pt set, spacing clean, taaki padhne mein maza aaye... 🔤👌",
    "Pandoc engine LaTeX math ko native Word equation blocks mein convert kar raha hai... ⚙️📐",
    "Page break protection lag rahi hai taaki question aadha idhar aadha udhar na kate... 🛡️📑",
    "Section dividers draw ho rahe hain, ekdum clean slate-gray finish ke saath... 📏🩶",
    "Homework title aur selected date top pe ekdum royal look mein baith rahe hain... 👑📅",
    "MCQ options ko 0.18 inch ka subtle indentation diya jaa raha hai... 📐📐",
    "Sub-parts a), b), c) ko 0.12 inch indent karke discipline mein khada kar rahe hain... 🚶‍♂️🚶‍♂️",
    "Logo perfectly header ke center mein align ho raha hai... 🖼️🎯",
    "Word XML elements ko polish kiya jaa raha hai, koi bug nahi bachega... 🧹💻",
    "Math formulas bilkul standard textbook jaise sharp dikhenge... 🏛️📐",
    "Bottom right corner pe dynamic 'Page X of Y' footer print ho raha hai... 📄🔢",
    "Spacing aisi adjust ho rahi hai ki page par na bheed ho, na khaali lage... ⚖️📄",
    "Paper ka first impression aisa hai ki bachhe dekhte hi bolenge: 'Yeh to easy lag raha hai'... 🎣😏",
    "Har question number aur option letter ko bold aur black kiya jaa raha hai... 🅱️🔲",
    "Reaction arrows ekdum straight aur aligned set ho chuke hain... ➡️📏",
    "Spacing after Option D ko 12 pt fix kar rahe hain for breathing room... 🌬️📏",
    "Document formatting ka level 100/100 touch kar raha hai... 💯🎨",
    "Aisa lag raha jaise kisi top institution ka assessment paper print ho raha ho... 🎓🏛️",
    "Professional styling in progress: Zero clutter, pure elegance... ✨📄",
    "Sabhi fractions aur superscripts ko microscope se check kiya jaa raha hai... 🔬🔢",
    "Word document ban gaya hai, ab PDF pipeline mein transfer ho raha hai... 🔄📑",
    "Virtual printing press start ho chuki hai... 🖨️⚡",
    "Document margins 0.8 inch pe perfectly calibrated... 📐📏",
    "Chemical notation packing: state symbols (s), (l), (g) ekdum neat... 🧪💧",
    "Paper ka look itna clean aur sharp hai ki frame karwane ka mann karega... 🖼️😎",
    "Keep_with_next logic ne questions ko page split se bacha liya hai... 🔒📄",
    "Header aur footer distance perfectly 0.35 aur 0.4 inches pe set... 📏📍",
    "Color palette: Deep Slate (#1E293B) and elegant slate lines... 🎨🖌️",
    "Formatting bugs ko jhadu maar ke bahar nikaal diya gaya hai... 🧹🐛",
    "PDF engine gears ghoom rahe hain, vector rendering chalu hai... ⚙️🖨️",
    "Har formula crystal clear vector curves mein convert ho raha hai... 📈✨",
    "No text distortion, zero equation blurriness... 💎📄",
    "Print preview dekha aur dil khush ho gaya... 🥰🖨️",
    "Bachhe sochenge teacher ne poori raat jaag ke itna sundar paper banaya hai... 🌙👨‍🏫",
    "Finishing touch chal rahi hai, jaise car ko showroom polish lagti hai... 🚗✨",
    "Document ready to print mode mein switch ho raha hai... 🖨️🏁",
    "Har question ka margin inspection pass ho chuka hai... 🛂📏",
    "Page count optimize ho chuka hai, koi unnecessary blank space nahi bacha... 🚫📄",
    "Vector typography locked, zoom karne pe bhi pixels nahi fatenge... 🔍👌",
    "Section titles ALL CAPS mein garv se khade hain... 🏛️🔠",
    "Assessment styling standard top-tier publication jaisi feel de rahi hai... 📜🏅",
    "PDF pages compile ho rahe hain, ek ke baad ek... 📚⚡",
    "Final proofreading AI inspector ne sign-off de diya hai... ✍️✅",
    "Look aisa hai ki homework dekh kar padhne ka mann kare (par solve na ho)... 📖😈",
    "Document formatting 99% ready, bas kuch hi pal bache hain... ⏳✨",
    "Styling wrap up! Ab launchpad pe final countdown... 🚀🎯"
]

TIER_6_MESSAGES = [
    "Boom! Jaadu complete. Ab mast evil smile ke saath homework bhej do! 🪄🎉",
    "Paper ready hai! Ab WhatsApp group pe bhej ke phone silent kar lijiye... 📱🔥",
    "Mission accomplished! Bachhon ke agle weekend ka poora bandobast ho gaya... 🏖️❌",
    "Tadaaa! Masterpiece ban gaya, ab print nikaalo aur distribution shuru karo! 🎓🚀",
    "Homework tayyar hai, ab class mein 'Pin drop silence' dekhne ke liye ready ho jaiye... 🤫📌",
    "PDF generate ho chuki hai, folder kholiye aur kamaal dekhiye! 📂✨",
    "Agli class mein viva aur homework submission ka darr shuru hone wala hai... 😨📝",
    "Kaam ho gaya boss! Bachhe yaad karenge ki kisse panga liya tha... 😈😎",
    "Sawaal itne kadak hain ki solutions check karte waqt aapko mazaa aayega... ☕📝",
    "Success! Aapka portable exam machine never disappoints... 🦾🏆",
    "Final PDF save ho gayi hai, ab bachhon ke excuses sunne ke liye taiyar rahein... 🎧🙄",
    "Magic finished! Ab print button dabaiye aur homework ka bomb phodiye... 💣🖨️",
    "Bachhon ka gaming session cancel karne wala legal document ready hai... 🎮🚫",
    "Super clean, super professional assessment ab aapke folder mein safe hai... 📁💎",
    "Ek aur successful batch ka dimaag kholne ka rasta saaf ho gaya... 🧠🔓",
    "Paper dekh kar principal sir/ma'am bhi bolenge: 'Outstanding work!'... 👏🏫",
    "File ready! Ab popcorn khate khate bachhon ke doubt messages padhiye... 🍿📱",
    "Done deal! Homework itna tagda hai ki parents bhi khush ho jayenge... 👨‍👩‍👧‍👦👍",
    "Ab class group mein send karke likhiye: 'Due tomorrow 8 AM sharp!'... ⏰💀",
    "Assessment paper download ready! Aapki mehnat bacha li humne... 🤝⏳",
    "All set! Ab dekhte hain class ka topper kitna score karta hai... 🧐💯",
    "Homework weapon ready for launch in 3... 2... 1... 🚀💥",
    "Ta-da! Beautiful Word doc aur sharp vector PDF aapke saamne hai... 📄✨",
    "Class mein shanti banaye rakhne ka 100% effective formula ready hai... 🧘‍♂️🏫",
    "File generate ho gayi, ab chai ka ek ghoont lijiye aur chill kijiye... ☕😌",
    "Students ke 'Sir questions bahut tough the' sunne ka countdown shuru... ⏳🗣️",
    "Zero error, zero delay! Machine ne apna jaadu dikha diya... ⚡🪄",
    "Complete! Ab students ko reality check dene ka waqt aa gaya hai... 🪞📉",
    "Done! Is homework ko complete karne mein bachhon ki night-out pakki hai... 🌙☕",
    "Assessment ready! Ab aap official tor par class ke sabse creative teacher hain... 👑👨‍🏫",
    "Folder mein dono files chamak rahi hain: .docx aur .pdf! 📂🌟",
    "Ek click mein homework tayyar, bachhe sochte reh jayenge ki kab banaya... 🧙‍♂️💨",
    "Bachhon ke phone notifications ab homework alerts se bharenge... 🔔📄",
    "Finish line crossed! PDF quality ekdum top-notch aayi hai... 🏁👌",
    "Agle do din tak bachhon ke paas time pass karne ka koi bahana nahi bachega... ⏳🛑",
    "Homework delivered like a pro! Time to celebrate... 🥂🎉",
    "Perfect compilation! Ek bhi formula ya option miss nahi hua... 🎯✅",
    "Ab WhatsApp group pe bhej ke airplane mode on kar lijiye... ✈️📵",
    "Kaam tamam! Ab students homework solve karenge aur aap aaram... 🛋️📖",
    "High voltage assessment generated successfully! ⚡📜",
    "Bachhe sochenge questions kaunse planet se mangwaye hain... 🪐👽",
    "Class test ho ya homework, aapka standard ab sabse upar hai... 📈🔝",
    "Done! Har ek minute vasool hone wala assessment tayyar... ⏱️✨",
    "Masterpiece complete! Word aur PDF dono aapki seva mein hazir hain... 🫡📄",
    "Ab class monitor ko call karke boliye: 'Sabko print out nikalne bolo'... 📞🖨️",
    "Homework ready, tensions shifted from teacher to students... 🔄💆‍♂️",
    "Kaam 100% done! Ab bas results ka intezar kijiye... 📊⏳",
    "File saved on desktop! Ekdum crystal clear vector quality ke sath... 🖥️✨",
    "Congratulations! Agla weekend bachhon ke liye study camp banne wala hai... ⛺📚",
    "Ta-da! Jaadu samapt, evil smile shuru! Ab distribution start karo! 🪄😈🎉"
]

# ---------------------------------------------------------
# CORE BACKEND FUNCTIONS & BULLETPROOF MATH GUARDS
# ---------------------------------------------------------
def sanitize_raw_json_string(raw_json_str: str) -> str:
    cleaned = raw_json_str.strip()
    if cleaned.startswith("```json"):
        cleaned = cleaned[7:]
    elif cleaned.startswith("```"):
        cleaned = cleaned[3:]
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3]
    cleaned = cleaned.strip()

    latex_cmd_regex = (
        r'(?<!\\)\\(?=(?:text|textbf|textrm|textit|tfrac|frac|times|tau|theta|to|'
        r'rightarrow|rightleftharpoons|right|rho|rangle|rm|nu|nabla|neq|not|neg|'
        r'beta|bar|bf|bm|begin|end|boldsymbol|bullet|Delta|alpha|gamma|sigma|'
        r'lambda|mu|pi|phi|omega|degree|circ|pm|mp|le|ge|approx|sim|equiv|'
        r'cdot|dots|cdots|left|sqrt|partial|sum|int|infty)\b)'
    )
    return re.sub(latex_cmd_regex, r'\\\\', cleaned)


def check_text_for_math_errors(text: str) -> tuple[bool, str]:
    if not text:
        return True, "Valid"

    # 1. Parity check
    dollar_count = len(re.findall(r'(?<!\\)\$', text))
    if dollar_count % 2 != 0:
        return False, f"Unbalanced '$' delimiter detected (found {dollar_count} '$' signs). Every '$' must be closed."

    # 2. Check for squished words
    if re.search(r'\b\d+(?:\.\d+)?\s*(?:molof|mLof|gof)\b', text, re.IGNORECASE):
        return False, "Fused unit detected (e.g., 'molof' or 'mLof'). Ensure units and words like 'of' are separate plain text."

    # 3. Check for chemical equation missing an arrow
    has_state_symbols = len(re.findall(r'\((?:s|l|g|aq)\)', text)) >= 2
    has_reactants_plus = '+' in text
    has_arrow = any(sym in text for sym in [r'\rightarrow', r'\rightleftharpoons', r'\to', '→', '⇌'])
    if has_state_symbols and has_reactants_plus and not has_arrow:
        return False, "Chemical reaction detected without a reaction arrow (\\rightarrow or \\rightleftharpoons)."

    return True, "Valid"


def verify_assessment_json(data: dict, req_mcq: int, req_saq: int, req_laq: int) -> tuple[bool, str]:
    if not isinstance(data, dict):
        return False, "Response root must be a valid JSON object."

    mcqs = data.get("multiple_choice_questions", [])
    saqs = data.get("short_answer_questions", [])
    laqs = data.get("long_answer_questions", [])

    if len(mcqs) != req_mcq:
        return False, f"Expected exactly {req_mcq} MCQs, got {len(mcqs)}."
    if len(saqs) != req_saq:
        return False, f"Expected exactly {req_saq} Short Answer questions, got {len(saqs)}."
    if len(laqs) != req_laq:
        return False, f"Expected exactly {req_laq} Long Answer questions, got {len(laqs)}."

    for i, q in enumerate(mcqs, 1):
        stem = str(q.get("stem", "")).strip()
        if not stem:
            return False, f"MCQ #{i} has an empty stem."
        ok, err = check_text_for_math_errors(stem)
        if not ok:
            return False, f"MCQ #{i} stem math error: {err}"

        opts = q.get("options", {})
        if not isinstance(opts, dict):
            return False, f"MCQ #{i} options must be a dictionary."
        for opt_key in ["A", "B", "C", "D"]:
            if opt_key not in opts or not str(opts[opt_key]).strip():
                return False, f"MCQ #{i} is missing option '{opt_key}'."
            ok, err = check_text_for_math_errors(str(opts[opt_key]))
            if not ok:
                return False, f"MCQ #{i} option {opt_key} math error: {err}"

    for i, q in enumerate(saqs, 1):
        stem = str(q.get("stem", "")).strip()
        if not stem:
            return False, f"Short Answer #{i} has an empty stem."
        ok, err = check_text_for_math_errors(stem)
        if not ok:
            return False, f"Short Answer #{i} stem math error: {err}"
        for s_idx, sub in enumerate(q.get("sub_parts", []), 1):
            ok, err = check_text_for_math_errors(str(sub))
            if not ok:
                return False, f"Short Answer #{i} sub-part #{s_idx} math error: {err}"

    for i, q in enumerate(laqs, 1):
        stem = str(q.get("stem", "")).strip()
        if not stem:
            return False, f"Long Answer #{i} has an empty stem."
        ok, err = check_text_for_math_errors(stem)
        if not ok:
            return False, f"Long Answer #{i} stem math error: {err}"
        sub_parts = q.get("sub_parts", [])
        if not sub_parts or not isinstance(sub_parts, list):
            return False, f"Long Answer #{i} must include sub_parts."
        for l_idx, sub in enumerate(sub_parts, 1):
            ok, err = check_text_for_math_errors(str(sub))
            if not ok:
                return False, f"Long Answer #{i} sub-part #{l_idx} math error: {err}"

    return True, "Valid"


def repair_math_syntax(text: str) -> str:
    """Bulletproof deterministic post-processor: un-fuses words, fixes enthalpy, inserts missing arrows, and preserves math blocks."""
    if not text:
        return ""

    # 1. Global word un-fusing
    text = re.sub(r'(\d+(?:\.\d+)?)\s*molof\s*([A-Za-z0-9_])', r'\1 mol of \2', text)
    text = re.sub(r'(\d+(?:\.\d+)?)\s*mLof\s*([A-Za-z0-9_])', r'\1 mL of \2', text)
    text = re.sub(r'(\d+(?:\.\d+)?)\s*gof\s*([A-Za-z0-9_])', r'\1 g of \2', text)
    text = re.sub(r'([a-zA-Z0-9_\)\}\]])and(\d+)', r'\1 and \2', text)
    text = re.sub(r'([a-zA-Z0-9_\)\}\]])and([A-Za-z])', r'\1 and \2', text)
    text = re.sub(r'\bmixed with(\d+)', r'mixed with \1', text)
    text = re.sub(r'\bwith(\d+)', r'with \1', text)
    text = re.sub(r'\)is\b', r') is ', text)
    text = re.sub(r'([a-zA-Z0-9\)])is mixed with', r'\1 is mixed with', text)
    text = re.sub(r'\$len~', r'', text)

    # 2. Fix malformed enthalpy notations (e.g., 'with H = -197^{-1}')
    text = re.sub(
        r'\b(?:with\s+)?(?:\\?Delta\s*)?H\s*=\s*(-?\d+(?:\.\d+)?)\s*(?:\^?\{?-?1\}?|kJ\s*mol\^?\{?-?1\}?|kJ\/mol)?\b',
        r'with $\\Delta H = \1\\text{ kJ mol}^{-1}$',
        text
    )

    # 3. Detect and fix missing reaction arrows in chemical reactions
    text = re.sub(
        r'(\((?:s|l|g|aq)\))\s+(?!(?:and|or|with|react|is|to|in|at|which|by|produce|from|determine)\b)(\d*[A-Z][a-zA-Z0-9_\(\)]*\(?[a-z]*\)?(?:\s*[\+]|\s*\((?:s|l|g|aq)\)))',
        r'\1 \\rightarrow \2',
        text
    )

    # 4. Protect double-dollar blocks ($$ ... $$) so they are never touched
    display_blocks = []
    def stash_display(m):
        display_blocks.append(m.group(0))
        return f"___DISPLAYBLOCK_{len(display_blocks)-1}___"
    text = re.sub(r'\$\$[^\$]+\$\$', stash_display, text)

    # 5. Safe inline math processing: NEVER strip blocks containing arrows or math operators
    def clean_inline(m):
        content = m.group(1).strip()
        # If it has chemistry arrows or standard math operators, KEEP IN MATH MODE
        if any(sym in content for sym in [r'\rightarrow', r'\rightleftharpoons', r'\Delta', r'\times', r'\frac', '=', '→', '⇌']):
            return f"${content}$"

        # Check for plain English prose mistakenly trapped in math
        stopwords = r'\b(is|are|was|were|in|into|absorbed|with|from|than|to|of|the|and|reacting|directly|rather|take|place|vessel|which|calculate|determine|explain|identify|predict|describe|mixed)\b'
        words_found = re.findall(stopwords, content, re.IGNORECASE)

        # Un-nest plain text prepositions from inside math
        if words_found or content.count(' ') >= 3:
            unwound = content
            unwound = re.sub(r'\b(is mixed with|mixed with|is|of|and|with)\b', r'$ \1 $', unwound)
            return unwound

        return f"${content}$"

    text = re.sub(r'(?<!\\)\$([^\$]+)\$', clean_inline, text)

    # Clean up any empty or duplicated $$ created by unwinding
    text = re.sub(r'\$\s*\$', ' ', text)

    # Restore display blocks
    for i, block in enumerate(display_blocks):
        text = text.replace(f"___DISPLAYBLOCK_{i}___", block)

    # 6. Ensure standalone reactions have proper math formatting for Pandoc
    def ensure_equation_math(line):
        if any(sym in line for sym in [r'\rightarrow', r'\rightleftharpoons', '→', '⇌']):
            if not line.strip().startswith('$') and not line.strip().endswith('$'):
                # Wrap bare reaction in math delimiters if not already inside
                line = re.sub(
                    r'((?:\d*[A-Z][a-zA-Z0-9_\(\)]*(?:\([s|l|g|aq]+\))?\s*[\+]\s*)+\d*[A-Z][a-zA-Z0-9_\(\)]*(?:\([s|l|g|aq]+\))?\s*(?:\\rightarrow|\\rightleftharpoons|→|⇌)\s*(?:\d*[A-Z][a-zA-Z0-9_\(\)]*(?:\([s|l|g|aq]+\))?\s*[\+]?\s*)+)',
                    r'$\1$',
                    line
                )
        return line

    text = "\n".join([ensure_equation_math(line) for line in text.splitlines()])

    # 7. Parity check: if an unclosed $ remains, close it
    dollar_count = len(re.findall(r'(?<!\\)\$', text))
    if dollar_count % 2 != 0:
        text = text + "$"

    return text


def get_available_models(client: genai.Client) -> list[str]:
    candidates = []
    try:
        models_pager = client.models.list()
        for m in models_pager:
            clean_name = m.name.replace("models/", "") if hasattr(m, "name") else ""
            if any(tag in clean_name.lower() for tag in ["flash", "pro"]) and not any(tag in clean_name.lower() for tag in ["vision", "tts", "clip", "transcribe", "image"]):
                candidates.append(clean_name)
    except Exception:
        pass

    combined = []
    for m in STATIC_MODELS_TO_TRY:
        if m not in combined:
            combined.append(m)

    for m in candidates:
        if m not in combined and "2.5-flash" not in m:
            combined.append(m)

    return combined


# ---------------------------------------------------------
# DYNAMIC AUSTRALIAN CURRICULUM PROMPT GENERATOR
# ---------------------------------------------------------
def construct_australian_system_prompt(year_level: str, req_mcq: int, req_saq: int, req_laq: int) -> str:
    if year_level == "Year 6":
        curriculum_context = """CURRICULUM CONTEXT: Australian Curriculum (ACARA v9.0) - Year 6 Primary Science.
- FOCUS: Observable properties, states of matter (solid, liquid, gas), reversible physical changes (melting, freezing, evaporating, condensing, dissolving) vs. irreversible chemical changes (rusting, burning, cooking, decaying).
- SCIENCE INQUIRY SKILLS: Fair testing principles (identifying independent variable [what we change], dependent variable [what we measure], and controlled variables [what we keep the same]). Reading simple data tables and drawing observational conclusions.
- PEDAGOGICAL LEVEL: Clear, engaging, accessible English. No complex molar mathematics."""
    elif year_level in ["Year 7", "Year 8"]:
        curriculum_context = """CURRICULUM CONTEXT: Australian Curriculum (ACARA v9.0) - Years 7-8 Junior Secondary Science (Chemical Sciences).
- FOCUS:
  * Particle model of matter (kinetic theory, arrangements and movement in solids, liquids, gases, gas pressure).
  * Pure substances vs. mixtures (solutions, suspensions, colloids, solutes, solvents).
  * Separation techniques (filtration, evaporation, simple distillation, chromatography, decanting, centrifuging).
  * Physical changes vs. chemical changes (color change, gas produced, precipitate formed, temperature change).
  * Introduction to elements, compounds, mixtures, and common element symbols. Simple word equations.
- SCIENCE INQUIRY SKILLS: Designing fair tests, writing hypotheses ("If... then..."), identifying laboratory errors (e.g. wet filter paper, improper thermometer placement), and interpreting line graphs.
- AUSTRALIAN REFERENCE BLUEPRINT: Model after authentic Australian National Assessment Science Inquiry items (ICAS Science, Big Science Competition)."""
    elif year_level in ["Year 9", "Year 10"]:
        curriculum_context = """CURRICULUM CONTEXT: Australian Curriculum (ACARA v9.0) - Years 9-10 Middle Secondary Science (Chemical Sciences).
- FOCUS:
  * Atomic structure (protons, neutrons, electrons, atomic number, mass number, electron configuration for elements 1-20).
  * Periodic Table organization (groups, periods, metals, non-metals, valence electrons).
  * Chemical reactions & conservation of mass: balancing equations, synthesis, combustion, precipitation, acid-base neutralisation, and acid-carbonate reactions.
  * Rates of reaction & collision theory: factors affecting rate (surface area, temperature, concentration, catalysts).
- SCIENCE INQUIRY SKILLS: Distinguishing between accuracy (calibration/systematic errors), reliability (consistency across concordant trials, spotting outliers), and validity (controlling external variables).
- AUSTRALIAN REFERENCE BLUEPRINT: Model after Australian state junior science exam formats (NSW Stage 5 Science Valid, Victorian Year 10 Science benchmarks)."""
    else:  # Year 11 (ATAR) or Year 12 (ATAR)
        curriculum_context = f"""CURRICULUM CONTEXT: Australian Senior Secondary Chemistry ({year_level} ATAR - NESA NSW Stage 6 / VCAA VCE / QCE Chemistry).
- FOCUS:
  * Quantitative chemistry: mole calculations ($n = m/M$, $c = n/V$), limiting reagents, percentage yield.
  * Gas stoichiometry under Australian Standard Laboratory Conditions (SLC: $25^\\circ\\text{{C}} / 298.15\\text{{ K}}$ and $100\\text{{ kPa}}$, where molar volume $V_m = 24.79\\text{{ L mol}}^{{-1}}$). (NEVER use US STP 22.4 L/mol).
  * Equilibrium and reversible systems ($K_c$, reaction quotient $Q$, Le Chatelier's principle, collision theory rates vs. yield compromise).
  * Acids and bases: Brønsted-Lowry theory, amphiprotic species, $K_a$, $pH$, titration curves, indicator selection.
  * Thermochemistry: $\\Delta H$, calorimetry calculation using $c = 4.18\\text{{ J g}}^{{-1}}\\text{{ K}}^{{-1}}$.
  * Organic chemistry & analysis: IUPAC nomenclature, functional groups, spectral interpretation.
- AUSTRALIAN ATAR COMMAND VERBS: Every short/long answer prompt must be anchored by official NESA/VCAA verbs:
  * "Explain in terms of...": Link cause and effect at the sub-microscopic/molecular level.
  * "Justify": Support conclusions by contrasting against an alternative scenario.
  * "Assess" / "Evaluate": Make evidence-based judgments (e.g. evaluating industrial compromises or experimental validity).
  * "Deduce": Derive an identity or trend from qualitative or quantitative data tables.
- WORKING SCIENTIFICALLY (EXPERIMENTAL RIGOR): At least 40% of short/long questions must involve laboratory investigations (titration error analysis, glassware rinsing procedures [water vs. titrant], outlier identification across repeated trials).
- AUSTRALIAN REFERENCE BLUEPRINT: Model after authentic NSW HSC Chemistry and Victorian VCE Chemistry Section I and Section II past examination papers."""

    return f"""You are a Chief Examination Item Writer and Senior Curriculum Assessor for the Australian Secondary School System.

Your objective is to craft an authentic, curriculum-aligned Australian Examination Paper for {year_level} based EXCLUSIVELY on the provided Chapter PDF.

================================================================================
1. STRICT CONTENT BOUNDARY & TOPIC ISOLATION (ZERO LEAKAGE)
================================================================================
- [DOCUMENT 1: CHAPTER PDF] defines the absolute syllabus boundary. Every question, scenario, substance, and calculation MUST originate 100% from this text.
- If a Past Examination Paper (PYQ) is provided [DOCUMENT 2], it serves SOLELY as an ARCHETYPE and PHRASING BLUEPRINT.
- NEVER copy past paper questions verbatim. Clone the inquiry structure, multi-part style, and question framing, but populate them exclusively with the concepts and data from Document 1.
- NEVER import unrelated topics from Document 2. If Document 2 contains topics outside Document 1, discard those topics completely.

================================================================================
2. PEDAGOGICAL & AUSTRALIAN CURRICULUM STANDARDS FOR {year_level.upper()}
================================================================================
{curriculum_context}

================================================================================
3. STRICT NEGATIVE CONSTRAINTS (FORMATTING & MARKS)
================================================================================
- DO NOT INCLUDE MARKS OR POINTS: Never write "[2 marks]", "(3 marks)", or total mark counts anywhere in the document.
- DO NOT GENERATE RESPONSE LINES OR WRITING SPACES: This is a pure question paper. Do not insert dotted lines, blank answer boxes, or empty response prompts.
- DO NOT INCLUDE ANSWER KEYS, RUBRICS, OR METADATA: Output only the questions and choices.
- ZERO META-LANGUAGE: Never write "refer to the text", "based on the chapter", or "according to the notes provided".

================================================================================
4. BULLETPROOF TYPOGRAPHY & LATEX RULES (STRICT NON-NEGOTIABLE)
================================================================================
1. NEVER BUNDLE ENGLISH WORDS OR UNITS INSIDE MATH MODE:
   * FORBIDDEN: "$50.0 mL of 1.00 mol L^{-1} Pb(NO_3)_2 is mixed with 75 mL of 0.500 mol L^{-1} KI$"
   * FORBIDDEN: "$2.0 mol of N_2(g) and 10 mol of H_2(g)$"
   * FORBIDDEN: "$10 mol of NH_3$"
   * CORRECT: "50.0 mL of 1.00 mol L$^{{-1}}$ $\\text{{Pb(NO}}_3)_2$ is mixed with 75 mL of 0.500 mol L$^{{-1}}$ $\\text{{KI}}$"
   * CORRECT: "When 2.0 mol of $\\text{{N}}_2(g)$ and 10 mol of $\\text{{H}}_2(g)$ react according to..."
   * Numbers, units (mL, g, mol, L), and English words (of, and, is, with, mixed, react) MUST BE PLAIN TEXT OUTSIDE DOLLAR SIGNS.

2. CHEMICAL REACTIONS MUST ALWAYS HAVE ARROWS:
   * Every reaction MUST use a standard LaTeX arrow: `\\rightarrow` or `\\rightleftharpoons`.
   * FORBIDDEN: "2SO_2(g) + O_2(g) 2SO_3(g)" (NEVER omit the arrow!)
   * WRAP FULL REACTIONS IN DOUBLE DOLLAR SIGNS:
     $$\\text{{2SO}}_2(g) + \\text{{O}}_2(g) \\rightleftharpoons \\text{{2SO}}_3(g)$$
     $$\\text{{CaCO}}_3(s) + 2\\text{{HCl}}(aq) \\rightarrow \\text{{CaCl}}_2(aq) + \\text{{H}}_2\\text{{O}}(l) + \\text{{CO}}_2(g)$$

3. ENTHALPY NOTATION:
   * ALWAYS include the Delta symbol (\\Delta) and full units: "$\\Delta H = -197\\text{{ kJ mol}}^{{-1}}$".
   * NEVER write "with H = -197^{{-1}}".

4. ESCAPING:
   * Double-escape all LaTeX backslashes inside JSON strings (e.g. `\\\\text{{}}`, `\\\\Delta`, `\\\\rightarrow`, `\\\\rightleftharpoons`).

================================================================================
5. REQUIRED JSON SCHEMA
================================================================================
Output STRICT, raw, parsable JSON matching this exact structure:
{{
  "multiple_choice_questions": [
    {{
      "number": 1,
      "stem": "Context-rich question stem with chemical species in $...$",
      "options": {{
        "A": "Plausible option or calculation trap",
        "B": "Plausible option or calculation trap",
        "C": "Plausible option or calculation trap",
        "D": "Plausible option or calculation trap"
      }}
    }}
  ],
  "short_answer_questions": [
    {{
      "number": 1,
      "stem": "Experimental stimulus or chemical context scenario",
      "sub_parts": [
        "a) Scaffolded prompt aligned to {year_level} inquiry expectations (no marks)",
        "b) Follow-up analytical or explanation prompt (no marks)"
      ]
    }}
  ],
  "long_answer_questions": [
    {{
      "number": 1,
      "stem": "Comprehensive investigation, data interpretation, or multi-step analysis",
      "sub_parts": [
        "a) Initial determination, identification, or quantitative step",
        "b) Mechanistic, molecular, or experimental justification",
        "c) Critical evaluation of variables, validity, error source, or compromise"
      ]
    }}
  ]
}}

EXACT COUNTS TO GENERATE:
- multiple_choice_questions: EXACTLY {req_mcq} questions. (Empty array [] if 0).
- short_answer_questions: EXACTLY {req_saq} questions. (Empty array [] if 0).
- long_answer_questions: EXACTLY {req_laq} questions. (Empty array [] if 0).
"""


def generate_and_verify_homework(chapter_pdf_path: str, pyq_pdf_path: str | None, api_key: str, year_level: str, difficulty: str, custom_keywords: str, req_mcq: int, req_saq: int, req_laq: int) -> dict:
    client = genai.Client(api_key=api_key)
    
    uploaded_chapter = client.files.upload(file=chapter_pdf_path)
    
    uploaded_pyq = None
    if pyq_pdf_path and os.path.exists(pyq_pdf_path):
        uploaded_pyq = client.files.upload(file=pyq_pdf_path)

    system_prompt = construct_australian_system_prompt(year_level, req_mcq, req_saq, req_laq)

    # 1. Map Cognitive Difficulty Profile
    if difficulty == "Easy":
        difficulty_instruction = (
            f"DIFFICULTY LEVEL: EASY ({year_level} Foundational):\n"
            "- Emphasize core definitions applied to familiar scenarios, direct observations, single-step questions, and straightforward relationships.\n"
            "- Distractors in MCQs should be distinct, avoiding complex multi-layer traps."
        )
    elif difficulty == "Medium":
        difficulty_instruction = (
            f"DIFFICULTY LEVEL: MEDIUM ({year_level} Standard Australian Curriculum Benchmark):\n"
            "- Balanced distribution: 30% foundational, 50% application & variable analysis, 20% evaluation.\n"
            "- Requires clear cause-and-effect reasoning and standard experimental inquiry interpretation."
        )
    else:  # Difficult
        difficulty_instruction = (
            f"DIFFICULTY LEVEL: DIFFICULT ({year_level} High-Discrimination / Band 6 / VCE 40+ / Extension):\n"
            "- 70% high-order thinking: novel experimental scenarios, unexpected data anomalies, identifying subtle systematic errors, multi-step problem solving.\n"
            "- Distractors in MCQs must target common high-achiever misconceptions and calculation traps."
        )

    # 2. Inject Optional Focus Keywords
    keywords_instruction = ""
    if custom_keywords.strip():
        keywords_instruction = (
            f"\nEDUCATOR FOCUS SUBTOPICS / KEYWORDS:\n"
            f"The educator has specifically requested priority emphasis on: '{custom_keywords.strip()}'.\n"
            f"Ensure at least 60% of questions directly target or contextualize these specific concepts.\n"
        )

    # 3. Assemble Conversation Contents
    user_prompt = f"""Generate the {year_level} assessment paper in strict accordance with the Australian Curriculum framework.

{difficulty_instruction}
{keywords_instruction}

IMPORTANT INSTRUCTIONS REGARDING ATTACHMENTS:
- The first attached document is [DOCUMENT 1: CHAPTER PDF]. ALL testable concepts, chemical formulas, and contexts MUST originate from Document 1.
"""
    if uploaded_pyq is not None:
        user_prompt += """- The second attached document is [DOCUMENT 2: PAST PAPER / PYQ]. Use Document 2 as an archetype and structural blueprint for authentic Australian questioning style. Do NOT copy questions verbatim. Do NOT import outside topics from Document 2.
"""
    else:
        user_prompt += f"""- No past paper was uploaded. Reference authentic Australian Curriculum examination blueprints (ICAS, Big Science, NESA HSC, VCE) appropriate for {year_level}.
"""

    user_prompt += f"""
CRITICAL REMINDER:
- NEVER bundle words like 'of', 'and', 'with', 'is' or numbers and units inside math mode.
- ALL chemical reactions MUST have arrows: \\rightarrow or \\rightleftharpoons.
- Enthalpy MUST be formatted as $\\Delta H = -197\\text{{ kJ mol}}^{{-1}}$.

Ensure EXACT question counts: {req_mcq} MCQs, {req_saq} Short Answer, and {req_laq} Long Answer questions.
Do NOT include marks, points, or dotted writing lines. Output valid JSON only."""

    try:
        model_errors = {}
        active_models = get_available_models(client)

        for model_name in active_models:
            conversation_contents = [uploaded_chapter]
            if uploaded_pyq is not None:
                conversation_contents.append(uploaded_pyq)
            conversation_contents.append(user_prompt)

            for _ in range(1, 3):
                try:
                    response = client.models.generate_content(
                        model=model_name,
                        contents=conversation_contents,
                        config=types.GenerateContentConfig(
                            system_instruction=system_prompt,
                            temperature=0.2,
                            response_mime_type="application/json"
                        )
                    )

                    sanitized = sanitize_raw_json_string(response.text)
                    parsed_json = json.loads(sanitized)
                    is_valid, validation_msg = verify_assessment_json(parsed_json, req_mcq, req_saq, req_laq)

                    if is_valid:
                        return parsed_json

                    conversation_contents.append(response.text)
                    conversation_contents.append(f"VALIDATION ERROR: {validation_msg}\nFix all math/unit issues and regenerate.")
                except Exception as err:
                    model_errors[model_name] = str(err)
                    break

        error_details = "\n".join([f"• {m}: {err}" for m, err in model_errors.items()])
        raise RuntimeError(f"All model fallbacks failed:\n{error_details}")
    finally:
        try:
            client.files.delete(name=uploaded_chapter.name)
        except Exception:
            pass
        if uploaded_pyq is not None:
            try:
                client.files.delete(name=uploaded_pyq.name)
            except Exception:
                pass


def json_to_markdown(data: dict, selected_date: str) -> str:
    md_lines = ["# Homework\n", f"(Date - {selected_date})\n"]
    mcqs = data.get("multiple_choice_questions", [])
    saqs = data.get("short_answer_questions", [])
    laqs = data.get("long_answer_questions", [])

    if mcqs:
        md_lines.append("## MULTIPLE CHOICE QUESTIONS\n")
        for q in mcqs:
            num = q.get("number", 1)
            stem = repair_math_syntax(str(q.get("stem", "")).strip())
            md_lines.append(f"**{num}.** {stem}\n")
            opts = q.get("options", {})
            for opt_key in ["A", "B", "C", "D"]:
                opt_val = repair_math_syntax(str(opts.get(opt_key, "")).strip())
                md_lines.append(f"**{opt_key})** {opt_val}\n")
            md_lines.append("")

    if saqs:
        md_lines.append("## SHORT ANSWER QUESTIONS\n")
        for q in saqs:
            num = q.get("number", 1)
            stem = repair_math_syntax(str(q.get("stem", "")).strip())
            md_lines.append(f"**{num}.** {stem}\n")
            for sub in q.get("sub_parts", []):
                sub_str = repair_math_syntax(str(sub).strip())
                sub_formatted = re.sub(r'^([a-d]\)|[ivx]+\))\s*', r'**\1** ', sub_str, flags=re.IGNORECASE)
                if not sub_formatted.startswith("**"):
                    sub_formatted = f"**-** {sub_formatted}"
                md_lines.append(f"{sub_formatted}\n")
            md_lines.append("")

    if laqs:
        md_lines.append("## LONG ANSWER QUESTIONS\n")
        for q in laqs:
            num = q.get("number", 1)
            stem = repair_math_syntax(str(q.get("stem", "")).strip())
            md_lines.append(f"**{num}.** {stem}\n")
            for sub in q.get("sub_parts", []):
                sub_str = repair_math_syntax(str(sub).strip())
                sub_formatted = re.sub(r'^([a-d]\)|[ivx]+\))\s*', r'**\1** ', sub_str, flags=re.IGNORECASE)
                if not sub_formatted.startswith("**"):
                    sub_formatted = f"**-** {sub_formatted}"
                md_lines.append(f"{sub_formatted}\n")
            md_lines.append("")

    return "\n".join(md_lines).strip()


def build_base_docx(markdown_text: str, output_docx_path: str):
    extra_args = ["--from=markdown+tex_math_dollars+tex_math_single_backslash", "--to=docx"]
    pypandoc.convert_text(source=markdown_text, to="docx", format="markdown", outputfile=output_docx_path, extra_args=extra_args)


def add_page_number_fields(footer_paragraph):
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.text = ""

    r1 = footer_paragraph.add_run("Page ")
    r1.font.name = "Arial"
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = RGBColor(100, 116, 139)
    fld1 = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w'))
    footer_paragraph._p.append(fld1)

    r2 = footer_paragraph.add_run(" of ")
    r2.font.name = "Arial"
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor(100, 116, 139)
    fld2 = parse_xml(r'<w:fldSimple %s w:instr="NUMPAGES"/>' % nsdecls('w'))
    footer_paragraph._p.append(fld2)


def add_section_divider_before(paragraph):
    divider_xml = parse_xml(
        r'<w:p %s>'
        r'  <w:pPr>'
        r'    <w:pBrd>'
        r'      <w:bottom w:val="single" w:sz="12" w:space="1" w:color="CBD5E1"/>'
        r'    </w:pBrd>'
        r'    <w:spacing w:before="240" w:after="160"/>'
        r'  </w:pPr>'
        r'  <w:r>'
        r'    <w:rPr><w:sz w:val="4"/></w:rPr>'
        r'    <w:t> </w:t>'
        r'  </w:r>'
        r'</w:p>' % nsdecls('w')
    )
    paragraph._p.addprevious(divider_xml)


def style_assessment_docx(docx_path: str, logo_img_path: str = "iwish_logo.jpg"):
    doc = Document(docx_path)
    section = doc.sections[0]

    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.4)

    # 1. Header Logo
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.text = ""

    if os.path.exists(logo_img_path):
        run_logo = header_p.add_run()
        run_logo.add_picture(logo_img_path, width=Inches(1.8))

    # 2. Dynamic Footer
    footer = section.footer
    footer_p = footer.paragraphs[0]
    add_page_number_fields(footer_p)

    # 3. Typography, Borders & Formatting
    section_keywords = ["MULTIPLE CHOICE QUESTIONS", "SHORT ANSWER QUESTIONS", "LONG ANSWER QUESTIONS"]
    is_first_section = True

    valid_paragraphs = [
        p for p in doc.paragraphs
        if p.text.strip() or p._p.xpath('.//m:oMathPara') or p._p.xpath('.//m:oMath')
    ]

    for idx, p in enumerate(valid_paragraphs):
        text = p.text.strip()
        clean_upper = text.upper().replace("#", "").strip()

        next_p = valid_paragraphs[idx + 1] if idx + 1 < len(valid_paragraphs) else None
        next_text = next_p.text.strip() if next_p else ""
        next_upper = next_text.upper().replace("#", "").strip()

        is_terminal = (
            next_p is None
            or any(s in next_upper for s in section_keywords)
            or bool(re.match(r'^\d+[\.\)]\s*', next_text))
        )

        if text.startswith("Homework") and len(text) < 15:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(14)
                r.font.bold = True
            continue

        if text.startswith("(Date -"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(14)
            p.paragraph_format.keep_with_next = True
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(71, 85, 105)
            continue

        matched_section = next((s for s in section_keywords if s in clean_upper), None)
        if matched_section:
            p.text = matched_section
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(13.5)
                run.font.bold = True
                run.font.underline = True
                run.font.color.rgb = RGBColor(30, 41, 59)

            if not is_first_section:
                add_section_divider_before(p)
                p.paragraph_format.space_before = Pt(6)
            else:
                p.paragraph_format.space_before = Pt(8)
                is_first_section = False

            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
            continue

        mcq_opt_match = re.match(r'^([A-D]\))\s*(.*)', text)
        if mcq_opt_match:
            opt_key = mcq_opt_match.group(1)
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.keep_together = True

            if opt_key.startswith(('D', 'd')) or is_terminal:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.keep_with_next = False
            else:
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.keep_with_next = True

            p.paragraph_format.line_spacing = 1.15
            continue

        sub_part_match = re.match(r'^([a-d]\)|[ivx]+\))\s*(.*)', text, re.IGNORECASE)
        if sub_part_match:
            p.paragraph_format.left_indent = Inches(0.12)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.keep_together = True

            if is_terminal:
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.keep_with_next = False
            else:
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
            continue

        q_match = re.match(r'^(\d+[\.\)])\s*(.*)', text)
        if q_match:
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.keep_together = True

            if is_terminal:
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.keep_with_next = False
            else:
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.keep_with_next = True
            continue

        has_math_block = bool(p._p.xpath('.//m:oMathPara') or (p._p.xpath('.//m:oMath') and len(text) < 5))
        if has_math_block:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.keep_together = True

            if is_terminal:
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.keep_with_next = False
            else:
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
            continue

        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.keep_together = True

        if is_terminal:
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = False
        else:
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True

    doc.save(docx_path)


def convert_docx_to_pdf(docx_path: str, output_pdf_path: str) -> bool:
    out_dir = os.path.dirname(output_pdf_path)
    try:
        cmd = ["libreoffice", "--headless", "--convert-to", "pdf", docx_path, "--outdir", out_dir]
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True, timeout=60)
        expected_pdf = os.path.join(out_dir, Path(docx_path).stem + ".pdf")
        if os.path.exists(expected_pdf):
            if expected_pdf != output_pdf_path:
                os.replace(expected_pdf, output_pdf_path)
            return True
    except Exception:
        pass
    return False


# ---------------------------------------------------------
# STREAMLIT USER INTERFACE & SESSION STATE MANAGEMENT
# ---------------------------------------------------------
if "assessment_data" not in st.session_state:
    st.session_state["assessment_data"] = None

st.title("✨ iWish Exam Paper Generator ✨")
st.markdown("Generate authentic Australian Curriculum & ATAR homework assessments from textbook PDFs.")

# Sidebar Settings
st.sidebar.header("⚙️ Configuration")
api_key = resolve_api_key()
if not api_key:
    api_key = st.sidebar.text_input("Enter Gemini API Key", type="password", help="Add to Streamlit Secrets to avoid typing.")
    if not api_key:
        st.sidebar.warning("API Key required to run the generator.")

year_level = st.sidebar.selectbox(
    "🎓 Grade / Year Level",
    options=[
        "Year 6",
        "Year 7",
        "Year 8",
        "Year 9",
        "Year 10",
        "Year 11 (ATAR)",
        "Year 12 (ATAR)"
    ],
    index=5,
    help="Select the Australian Curriculum or Senior ATAR level."
)

selected_date = st.sidebar.date_input("Homework Date", value=datetime.date.today()).strftime("%d.%m.%Y")

difficulty = st.sidebar.select_slider(
    "📊 Difficulty Level",
    options=["Easy", "Medium", "Difficult"],
    value="Medium",
    help="Adjusts foundational vs. high-order thinking question distribution."
)

st.sidebar.subheader("Question Breakdown (0 to 10)")
req_mcq = st.sidebar.slider("Multiple Choice Questions", min_value=0, max_value=10, value=5)
req_saq = st.sidebar.slider("Short Answer Questions", min_value=0, max_value=10, value=5)
req_laq = st.sidebar.slider("Long Answer Questions", min_value=0, max_value=10, value=0)

if req_mcq + req_saq + req_laq == 0:
    st.sidebar.error("Select at least 1 question to generate.")

enable_keywords = st.sidebar.checkbox("🎯 Target Specific Subtopics / Keywords", value=False)
custom_keywords = ""
if enable_keywords:
    custom_keywords = st.sidebar.text_area(
        "Enter keywords or subtopics to prioritize:",
        placeholder="e.g., Buffer action, Ka calculations, phenolphthalein endpoint, titration errors",
        help="Gemini will allocate at least 60% of questions to these concepts."
    )

st.subheader("📁 Document Inputs")
col_upload1, col_upload2 = st.columns(2)

with col_upload1:
    uploaded_chapter = st.file_uploader(
        "📘 Upload Chapter PDF (Mandatory Syllabus Content)",
        type=["pdf"],
        help="The questions will strictly cover ONLY concepts in this document."
    )

with col_upload2:
    uploaded_pyq = st.file_uploader(
        "📝 Upload Past Paper / PYQ (Optional Style Blueprint)",
        type=["pdf"],
        help="Used purely as a style & question structure archetype. Concepts from other chapters in this file are ignored."
    )

col_action1, col_action2 = st.columns([3, 1])
with col_action1:
    generate_clicked = st.button("🪄 Generate Assessment Paper", type="primary", use_container_width=True)
with col_action2:
    if st.session_state["assessment_data"] is not None:
        if st.button("🔄 Clear / New Exam", use_container_width=True):
            st.session_state["assessment_data"] = None
            st.rerun()

if generate_clicked:
    if not api_key:
        st.error("Please enter a valid Gemini API Key in the sidebar or deploy secrets.")
        st.stop()
    if req_mcq + req_saq + req_laq == 0:
        st.error("Please select at least one question type.")
        st.stop()
    if not uploaded_chapter:
        st.error("Please upload the Chapter PDF before generating.")
        st.stop()

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_chapter_pdf = os.path.join(temp_dir, uploaded_chapter.name)
        with open(temp_chapter_pdf, "wb") as f:
            f.write(uploaded_chapter.getbuffer())

        temp_pyq_pdf = None
        if uploaded_pyq:
            temp_pyq_pdf = os.path.join(temp_dir, uploaded_pyq.name)
            with open(temp_pyq_pdf, "wb") as f:
                f.write(uploaded_pyq.getbuffer())

        base_stem = Path(uploaded_chapter.name).stem
        docx_output_path = os.path.join(temp_dir, f"{base_stem}_Assessment.docx")
        pdf_output_path = os.path.join(temp_dir, f"{base_stem}_Assessment.pdf")

        progress_bar = st.progress(0)
        status_box = st.empty()

        try:
            # Tier 1 (0% -> 20%)
            progress_bar.progress(10)
            status_box.info(random.choice(TIER_1_MESSAGES))
            time.sleep(1.0)

            # Tier 2 (20% -> 40%)
            progress_bar.progress(25)
            status_box.info(random.choice(TIER_2_MESSAGES))

            # Backend AI Generation
            assessment_json = generate_and_verify_homework(
                chapter_pdf_path=temp_chapter_pdf,
                pyq_pdf_path=temp_pyq_pdf,
                api_key=api_key,
                year_level=year_level,
                difficulty=difficulty,
                custom_keywords=custom_keywords,
                req_mcq=req_mcq,
                req_saq=req_saq,
                req_laq=req_laq
            )

            # Tier 3 (40% -> 60%)
            progress_bar.progress(50)
            status_box.info(random.choice(TIER_3_MESSAGES))
            time.sleep(1.0)

            markdown_content = json_to_markdown(assessment_json, selected_date)

            # Tier 4 (60% -> 80%)
            progress_bar.progress(68)
            status_box.info(random.choice(TIER_4_MESSAGES))

            build_base_docx(markdown_content, docx_output_path)

            # Tier 5 (80% -> 95%)
            progress_bar.progress(85)
            status_box.info(random.choice(TIER_5_MESSAGES))

            style_assessment_docx(docx_output_path, logo_img_path="iwish_logo.jpg")

            pdf_success = convert_docx_to_pdf(docx_output_path, pdf_output_path)

            # Tier 6 (100%)
            progress_bar.progress(100)
            status_box.success(random.choice(TIER_6_MESSAGES))
            time.sleep(0.8)

            with open(docx_output_path, "rb") as f_docx:
                docx_bytes = f_docx.read()

            pdf_bytes = None
            if pdf_success and os.path.exists(pdf_output_path):
                with open(pdf_output_path, "rb") as f_pdf:
                    pdf_bytes = f_pdf.read()

            st.session_state["assessment_data"] = {
                "docx_bytes": docx_bytes,
                "pdf_bytes": pdf_bytes,
                "base_stem": base_stem
            }
            status_box.empty()
            progress_bar.empty()
            st.rerun()

        except Exception as e:
            status_box.empty()
            progress_bar.empty()
            st.error(f"Generation error: {e}")

# ---------------------------------------------------------
# PERSISTENT DOWNLOADS & PDF PREVIEW
# ---------------------------------------------------------
if st.session_state["assessment_data"] is not None:
    data = st.session_state["assessment_data"]

    st.success("🎉 Assessment generation complete! You can download both files below.")

    col1, col2 = st.columns(2)
    with col1:
        st.download_button(
            label="📥 Download Word Assessment (.docx)",
            data=data["docx_bytes"],
            file_name=f"{data['base_stem']}_Assessment.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    with col2:
        if data["pdf_bytes"]:
            st.download_button(
                label="📥 Download Vector PDF (.pdf)",
                data=data["pdf_bytes"],
                file_name=f"{data['base_stem']}_Assessment.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        else:
            st.info("PDF engine conversion unavailable. Download DOCX to print/save as PDF.")

    if data["pdf_bytes"]:
        st.markdown("---")
        st.subheader("📄 In-Browser Assessment Preview")
        base64_pdf = base64.b64encode(data["pdf_bytes"]).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="850" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)
